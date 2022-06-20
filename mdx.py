from copy import copy
from pathlib import Path
import re
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor, Pt, Cm
import sys

STYLE_CODE = "Code"
CLI_HELP = "Usage: mdx [in] [out]\n\n  Seemless markdown to docx converter\n\nArguments:\n  --andy    Alternate high-clarity document format"


class Context:
    """Contextual information for compartmentalised converting"""

    line = 0
    heading = None
    italic = False
    bold = False
    underline = False
    strikethrough = False

    def no_spacing(self) -> bool:
        """Checks if elements should have spacing within the current section"""
        if self.heading is None:
            return False
        return self.heading.text.lower() in ["bibliography", "references"]

    def next_line(self):
        """Skips to the next line"""
        self.line += 1
        self.char = 0
        self.italic = False
        self.bold = False
        self.underline = False
        self.strikethrough = False

    def flip_italic(self):
        """Flips italic value"""
        self.italic = not self.italic

    def flip_bold(self):
        """Flips bold value"""
        self.bold = not self.bold


class Heading:
    """Heading section inside document"""

    def __init__(self, text: str, level: int) -> None:
        self.text = text
        self.level = level

    def _md(line: str):
        # Parse number of # for level
        level = 0
        while line[level] == "#":
            level += 1
        # Get and clean text
        text = line[level:].lstrip()
        return Heading(text, level)

    def _docx(self, docx_doc: docx.Document):
        docx_para = docx_doc.add_heading(self.text, self.level)
        # print(docx_para._element.xml)
        # docx_para.insert(0, etree.XML("<hi />"))


class Run:
    """Run of text with styling located inside a paragraph"""

    def __init__(self, ctx: Context, text: str, **kwargs):
        # Check that run is a string; python doesn't have strong typing sadly
        if type(text) != str:
            raise Exception("Make sure this run is a string, this is a common mistake")
        # Create tuns
        self.ctx = ctx
        self.text = text
        self.link = None
        self.link_external = None
        self.image = False
        if "link" in kwargs:
            self.link_external = kwargs["link"][1]

    def _docx(self, docx_para: docx.text.paragraph.Paragraph) -> docx.text.run.Run:
        # Act different if it's a link
        if self.link is not None:
            return _add_link(docx_para, self.link, self.text, self.link_external)
        # Add plain run text
        docx_run = docx_para.add_run(self.text)
        # Add relevant styles
        if self.ctx.bold:
            docx_run.bold = True
        if self.ctx.italic:
            docx_run.italic = True
        if self.ctx.underline:
            docx_run.underline = True
        if self.ctx.strikethrough:
            docx_run.strikethrough = True
        return docx_run


class Paragraph:
    """Paragraph consisting of many runs of text"""

    def __init__(self, ctx: Context, runs: list = []):
        self.ctx = ctx
        self.runs = runs

    def append(self, run: Run):
        """Appends new run to paragraph"""
        self.runs.append(run)

    @staticmethod
    def _md(ctx: Context, line: str):
        # Parse through runs
        runs = []
        ind = 0
        flipflop = False
        buf = ""
        add = True

        # Go through each character
        while ind < len(line):
            # Get character
            c = line[ind]

            # Bold/italics
            if c == "*":
                # Calculate flipflop
                add = flipflop
                if flipflop:
                    flipflop = False
                    continue
                # Finish existing buffer
                runs.append(Run(copy(ctx), buf))
                buf = ""
                # Get star length
                stars = len(line[ind:]) - len(line[ind:].lstrip("*"))
                # Italics if theres a non-even amount
                if stars % 2 == 1:
                    ctx.flip_italic()
                # Bold if theres two or more
                if stars > 1:
                    ctx.flip_bold()
                ind += stars - 1

            # Link (external or internal)
            match = re.search(
                r"^\[[^\]]*\]\((#[a-zA-Z0-9-]*|[-a-zA-Z0-9@:%_\+.~#?&//=]{2,256}\.[a-z]{2,4}\b(\/[-a-zA-Z0-9@:%_\+.~#?&//=]*)?)\)",
                line[ind:],
            )
            if match:
                # Finish existing buffer and skip link
                runs.append(Run(copy(ctx), buf))
                buf = ""
                add = False
                ind += len(match.group(0)) - 1
                # Parse components
                splitted = match.group(0).split("](", 1)
                link = splitted[1][:-1].strip()
                # Add link
                if link.startswith("#"):
                    # Internal link
                    text = splitted[0][1:]
                    runs.append(Run(copy(ctx), text, link=(link[1:], False)))
                else:
                    # External link
                    text = splitted[0][1:]  # TODO: parse markdown rather than raw text
                    # TODO: include local uris as an automatic appendix :)
                    runs.append(Run(copy(ctx), text, link=(link, True)))

            # Add to ind/buf
            if add:
                buf += c
            else:
                add = True
            ind += 1

        # Create paragraph and return
        runs.append(Run(copy(ctx), buf))
        return Paragraph(ctx, runs)

    def _docx(self, docx_doc: docx.Document) -> docx.text.paragraph.Paragraph:
        # Add empty paragraph
        docx_para = docx_doc.add_paragraph()
        # Make no-spaced if defined
        if self.ctx.no_spacing():
            docx_para.style = "No Spacing"
        # Add runs to paragraph
        for run in self.runs:
            run._docx(docx_para)
        return docx_para


class Codeblock:
    """Codeblock containing language and monospaced code"""

    def __init__(self, lines: list, lang: str = None, heading_after: bool = False):
        self.lines = lines
        self.lang = lang  # TODO: use somewhere in docx
        self.heading_after = heading_after

    @staticmethod
    def _md(lines: list) -> tuple:
        # Get language after ``` designator
        lang = (
            lines[0].lstrip()[3:].lstrip()
        )  # first `lstrip()` used in document parsing
        lang = lang if lang != "" else None

        # Read lines
        heading_after = False
        code = []
        for ind, line in enumerate(lines[1:]):
            if line.lstrip() == "```":
                # Check if there's a heading afterwards
                if len(lines[1:]) - 1 > ind and lines[ind + 2].lstrip().startswith("#"):
                    heading_after = True
                # Stop codeblock
                break
            else:
                code.append(line)

        # Get skip
        skip = len(code) + 1
        return (Codeblock(code, lang, heading_after), skip)

    def _docx(self, docx_doc: docx.Document):
        # Calculate justification for lines
        just = len(str(len(self.lines)))
        # Add lines
        for ind, line in enumerate(self.lines):
            # Figure out line number
            num = str(ind + 1).rjust(just)
            # Add new paragraph with code style
            docx_para = docx_doc.add_paragraph()
            docx_para.style = STYLE_CODE
            # Add line number with italics
            docx_run = docx_para.add_run(num)
            docx_run.font.italic = True
            # Add actual code
            docx_para.add_run(" " + line)

        # Add small codeblock line for formatting if there's not a heading afterwards
        if not self.heading_after:
            docx_para = docx_doc.add_paragraph()
            docx_para.style = STYLE_CODE


class Quote(Paragraph):
    """Quote of something in it's own style"""

    @staticmethod
    def _md(ctx: Context, line: str):
        # Level info
        level, line = _level_info(line)
        # Clean line from `>` starter
        line = line[1:].lstrip()
        # Parse via inheritance and convert
        para = super(Quote, Quote)._md(
            ctx, line
        )  # BODGE: python doesn't like staticmethod and inheritance
        quote = Quote(ctx, para.runs)
        # Set levelling
        quote.level = level
        return quote

    def _docx(self, docx_doc: docx.Document) -> docx.text.paragraph.Paragraph:
        # Get inherited generated paragraph
        para = super()._docx(docx_doc)
        # Reset to quote styling
        para.style = "Quote"
        INDENT = 0.75
        para.paragraph_format.left_indent = Cm(INDENT * self.level + 1)
        para.paragraph_format.right_indent = Cm(INDENT)
        return para


class PointBullet(Paragraph):
    """Bullet point with content inside of it"""

    @staticmethod
    def _md(ctx: Context, line: str):
        # Level info
        level, line = _level_info(line)
        # Clean line from `-` starter
        line = line[1:].lstrip()
        # Parse via inheritance and convert
        para = super(PointBullet, PointBullet)._md(
            ctx, line
        )  # BODGE: python doesn't like staticmethod and inheritance
        bullet = PointBullet(ctx, para.runs)
        # Set levelling
        bullet.level = level
        return bullet

    def _docx(self, docx_doc: docx.Document) -> docx.text.paragraph.Paragraph:
        # Get inherited generated paragraph
        docx_para = super()._docx(docx_doc)
        # Set bullet style according to level
        docx_para.style = (
            "List Bullet" if self.level == 0 else f"List Bullet {self.level+1}"
        )
        return docx_para


class PointNumbered(Paragraph):
    """Numbered point with content inside of it"""

    @staticmethod
    def _md(ctx: Context, line: str):
        # Level info
        level, line = _level_info(line)
        # Get number and clean
        splitted = line.split(".", 1)
        num = int(splitted[0])
        line = splitted[1].lstrip()
        # Parse via inheritance and convert
        para = super(PointNumbered, PointNumbered)._md(
            ctx, line
        )  # BODGE: python doesn't like staticmethod and inheritance
        numbered = PointNumbered(ctx, para.runs)
        # Set info
        numbered.level = level
        numbered.num = num
        return numbered

    def _docx(self, docx_doc: docx.Document) -> docx.text.paragraph.Paragraph:
        # TODO: use something like "start at self.num" so markdown starting at like `20.` can be used, it fucks up otherwise
        # Get inherited generated paragraph
        docx_para = super()._docx(docx_doc)
        # Set bullet style according to level
        docx_para.style = (
            "List Number" if self.level == 0 else f"List Number {self.level+1}"
        )
        return docx_para


class Image:
    """Image with some optional caption text"""

    def __init__(self, ctx: Context, link: str, caption: Paragraph = None) -> None:
        self.ctx = ctx
        self.link = link
        self.caption = caption

    @staticmethod
    def _md(ctx: Context, matched: str):
        splitted = matched.split("](")
        caption = splitted[0][2:].strip()
        link = splitted[1][:-1].strip()
        return Image(ctx, link, Paragraph._md(ctx, caption) if caption != "" else None)

    def _docx(self, docx_doc: docx.Document) -> list[docx.text.paragraph.Paragraph]:
        # Insert image
        # TODO: width/height adjustment algorithm
        docx_para_image = docx_doc.add_paragraph()
        docx_run = docx_para_image.add_run()
        docx_run.add_picture(self.link, height=Cm(10))

        # Add caption
        if self.caption:
            docx_para_caption = self.caption._docx(docx_doc)
            docx_para_caption.style = "Caption"

            return [docx_para_image, docx_para_caption]
        return [docx_para_image]


class Document:
    """High-level document abstractions for conversion"""

    # Default fonts (changes via andy setting)
    font_heading = "IBM Plex Sans"
    font_body = "IBM Plex Serif"
    font_code = "IBM Plex Mono"

    def __init__(self, md: str, andy: bool = False):
        # Components
        self.elements = []
        self.title = None
        self.subtitle = None
        self.ctx = Context()

        # Set andy format
        self.andy = andy
        if self.andy:
            self.font_heading = "Arial"
            self.font_body = "Arial"
            self.font_code = "Lucida Sans Typewriter"

        # Remove toc and clear up lines
        lines_raw = _rm_toc(md)
        lines = []
        for line in lines_raw:
            lines.append(line.rstrip())

        # Metadata
        if len(lines) > 1 and lines[0] == "---":
            # Go over lines in metadata
            skip = 0
            for ind, line in enumerate(lines[1:]):
                # Stop metadata if it's ended
                if line == "---":
                    skip = ind + 1
                    break
                # Split at `:` token
                splitted = line.split(":", 1)
                # Go to next line if its invalid
                if len(splitted) != 2:
                    continue
                # Clean left and right sections
                left = splitted[0].lstrip().lower()
                right = splitted[1].lstrip()
                # Match left section
                if left == "title":
                    self.title = right
                elif left == "subtitle":
                    self.subtitle = right
            # Skip to end of metadata if there was an open and close tag
            if skip != 0:
                lines = lines[1 + skip :]

        # Parse through lines
        while self.ctx.line < len(lines):
            # Get line
            line = lines[self.ctx.line]
            stripped = line.lstrip()
            # Check start
            if stripped.startswith("#"):
                # Heading
                heading = Heading._md(stripped)
                self.elements.append(heading)
                self.ctx.heading = heading
                # Check if heading defines references
            elif stripped.startswith("```"):
                # Codeblock
                codeblock, skip = Codeblock._md(lines[self.ctx.line :])
                self.ctx.line += skip
                self.elements.append(codeblock)
            elif stripped.startswith(">"):
                # Quote
                self.elements.append(Quote._md(copy(self.ctx), line))
            elif stripped.startswith("-"):
                # Bullet point
                self.elements.append(PointBullet._md(copy(self.ctx), line))
            elif match := re.search(
                r"^!?\[[^\]]*\]\((#[a-zA-Z0-9-]*|[-a-zA-Z0-9@:%_\+.~#?&//=]{2,256}\.[a-z]{2,4}\b(\/[-a-zA-Z0-9@:%_\+.~#?&//=]*)?)\)",
                line,
            ):
                # Image
                self.elements.append(Image._md(copy(self.ctx), match.group(0)))
            # Check others
            else:
                # Numbered point
                try:
                    if "." not in stripped:
                        raise Exception()
                    int(stripped.split(".", 1)[0])
                    self.elements.append(PointNumbered._md(copy(self.ctx), line))
                # Paragraph
                except:
                    if (
                        # Non-sensitive typical empty lines
                        (not self.ctx.no_spacing() and line == "")
                        # Sensitive but last line was title
                        or (
                            self.ctx.no_spacing()
                            and lines[self.ctx.line - 1].lstrip().startswith("#")
                        )
                        # Sensitive but next line is title
                        or (
                            self.ctx.no_spacing()
                            and len(lines) > self.ctx.line + 1
                            and lines[self.ctx.line + 1].lstrip().startswith("#")
                        )
                    ):
                        # Skip empty line
                        self.ctx.next_line()
                        continue
                    self.elements.append(Paragraph._md(copy(self.ctx), stripped))

            # Move to next line
            self.ctx.next_line()

    def save(self, path: Path):
        """Saves document to `path` provided"""
        # Create docx file
        docx_doc = docx.Document()

        # New styles
        style_codeblock = docx_doc.styles.add_style(STYLE_CODE, WD_STYLE_TYPE.PARAGRAPH)

        # Add title/subtitle
        if self.title or self.subtitle:
            # Create empty lines before title
            for _ in range(4):
                para = Paragraph(copy(self.ctx), [Run(copy(self.ctx), "")])
                para._docx(docx_doc)

            # Add title
            if self.title:
                docx_para = docx_doc.add_heading(self.title, 0)
            # Add subtitle
            if self.subtitle:
                docx_para = Paragraph(
                    copy(self.ctx), [Run(copy(self.ctx), self.subtitle)]
                )._docx(docx_doc)
                docx_para.style = "Subtitle"

            # Page break
            docx_para = docx_doc.add_paragraph()
            docx_run = docx_para.add_run()
            docx_run.add_break(WD_BREAK.PAGE)

        # Add elements
        for element in self.elements:
            element._docx(docx_doc)

        # Replace all fonts with body font by default
        for style in docx_doc.styles:
            if hasattr(style, "font"):
                style.font.name = self.font_body

        # Styling for title
        style_title = docx_doc.styles["Title"]
        _style_title_border(style_title)
        style_title.font.name = self.font_heading
        style_title.font.size = Pt(26)
        style_title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style_title.paragraph_format.space_after = Pt(3)
        style_title.paragraph_format.alignment = 1

        # Styling for subtitle
        style_subtitle = docx_doc.styles["Subtitle"]
        style_subtitle.font.name = self.font_heading
        style_subtitle.font.size = Pt(14)
        style_subtitle.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style_subtitle.font.italic = False
        style_subtitle.paragraph_format.alignment = 1

        # Styling for headings
        for h in range(1, 9):
            style_heading = docx_doc.styles[f"Heading {h}"]
            style_heading.font.name = self.font_heading
            style_heading.font.bold = self.andy
            style_heading.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            # Per-level styling
            if h == 1:
                style_heading.font.size = Pt(22)
                style_heading.paragraph_format.space_after = Pt(2)
            elif h == 2:
                style_heading.font.size = Pt(17)
            elif h <= 4:
                style_heading.font.size = Pt(13)
            # Italics for small headings
            if h > 3:
                style_heading.font.italic = True

        # Styling for paragraphs
        style_paragraph = docx_doc.styles["Normal"]
        if not self.andy:
            style_paragraph.paragraph_format.alignment = 3

        # Styling for no spacing
        style_nospace = docx_doc.styles["No Spacing"]
        style_nospace.paragraph_format.alignment = (
            0  # shouldn't be justified but might inherit from paragraph
        )

        # Andy styling for paragraphs and no spacing
        if self.andy:
            style_paragraph.paragraph_format.line_spacing = 1.5
            for style in [style_paragraph, style_nospace]:
                style.font.size = Pt(12)

        # Styling for codeblocks
        style_codeblock.font.name = self.font_code
        style_codeblock.paragraph_format.space_after = Pt(0)
        style_codeblock.paragraph_format.line_spacing = 1
        style_codeblock.paragraph_format.alignment = 0

        # TODO: new "Link" run styling

        # Use docx's vanilla save
        docx_doc.save(path)


def _style_title_border(style_title):
    """Removes border style on title which is set by python-docx by default.
    This is a hack because there's no programmatic way to do this as of writing"""
    el = style_title._element
    el.remove(el.xpath("w:pPr")[0])


def _level_info(line: str) -> tuple:
    """Figures out level information and returns it and the line without spacing"""
    stripped = line.lstrip()
    num = len(line) - len(stripped)
    level = int(num / 2)
    return (level, stripped)


def _err_exit(msg: str):
    """Prints error message to console and exits program, used for command-line"""
    print(f"{CLI_HELP}\n\nError: {msg}", file=sys.stderr)
    sys.exit(1)


def _rm_toc(md: str) -> list:
    """Removes first table of contents section from a markdown string, returning list of lines"""
    # Don't check if there isn't one
    check = md.lower()
    if "table of contents" not in check and "contents" not in check:
        return md.splitlines()
    # Parse through
    in_toc = False
    removed_toc = False
    keep = []
    for line in md.splitlines():
        clean = line.lstrip()
        # Title, so either start/end toc removal
        if clean.startswith("#") and not removed_toc:
            # Stop removing toc
            if in_toc:
                in_toc = False
                keep.append(line)
                continue
            # Start removing toc
            title = clean.lstrip("#").strip().lower()
            if title in ["table of contents", "contents"]:
                in_toc = True
            else:
                keep.append(line)
        # Add like normal
        elif not in_toc:
            keep.append(line)
    return keep


def _add_link(paragraph: str, link: str, text: str, external: bool):
    """Places an internal or external link within a paragraph object"""

    # Create the w:hyperlink tag
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    # Set where it links to
    if external:
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(
            link, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        # External relationship value
        hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)
    else:
        # Internal anchor value
        hyperlink.set(docx.oxml.shared.qn("w:anchor"), link)
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement("w:r")
    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    # Add link styling
    rStyle = docx.oxml.shared.OxmlElement("w:pStyle")
    rStyle.set(docx.oxml.shared.qn("w:val"), "Link")
    rPr.append(rStyle)
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # Add to paragraph
    paragraph._p.append(hyperlink)
    return hyperlink


# Command-line
if __name__ == "__main__":
    # Get and clean arguments
    args = sys.argv[1:]
    # Make sure theres at least an input and output or show help
    if len(args) == 0:
        _err_exit("Please provide [in]")
    elif "--help" in args[2:]:
        print(CLI_HELP)
        sys.exit(0)
    # Get andy setting
    andy = "--andy" in args[2:]
    # Get markdown from file
    md = ""
    try:
        with open(args[0], "r") as file:
            md = file.read()
    except Exception as e:
        _err_exit(f"Invalid file, {e}")
    # Create and save document to defined parts
    # try:
    Document(md, andy).save(args[1])
    # except Exception as e:
    #     _err_exit(f"Couldn't convert document; {e}")  # TODO: better errors
